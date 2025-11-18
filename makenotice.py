#!/usr/bin/env python3
"""
Notice Generator CLI Tool
Generates bank notices from Excel file and Word template
Based on RBI IFSC code structure:
- First 4 characters: Bank code (alphabetic)
- 5th character: Always '0' (reserved for future use)
- Last 6 characters: Branch code (numeric/alphanumeric)
"""

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
import argparse
import os
from collections import defaultdict

# Comprehensive IFSC to Bank Name mapping based on RBI specifications
# First 4 characters of IFSC identify the bank
IFSC_BANK_MAP = {
    # Public Sector Banks
    'SBIN': 'STATE BANK OF INDIA',
    'ALLA': 'ALLAHABAD BANK',
    'ANDB': 'ANDHRA BANK',
    'BARB': 'BANK OF BARODA',
    'BKID': 'BANK OF INDIA',
    'MAHB': 'BANK OF MAHARASHTRA',
    'CNRB': 'CANARA BANK',
    'CBIN': 'CENTRAL BANK OF INDIA',
    'CORP': 'CORPORATION BANK',
    'BKDN': 'DENA BANK',
    'IBKL': 'IDBI BANK',
    'IDIB': 'INDIAN BANK',
    'IOBA': 'INDIAN OVERSEAS BANK',
    'ORBC': 'ORIENTAL BANK OF COMMERCE',
    'PSIB': 'PUNJAB & SIND BANK',
    'PUNB': 'PUNJAB NATIONAL BANK',
    'SBBJ': 'STATE BANK OF BIKANER & JAIPUR',
    'SBHY': 'STATE BANK OF HYDERABAD',
    'SBMY': 'STATE BANK OF MYSORE',
    'STBP': 'STATE BANK OF PATIALA',
    'SBTR': 'STATE BANK OF TRAVANCORE',
    'SYNB': 'SYNDICATE BANK',
    'UCBA': 'UCO BANK',
    'UBIN': 'UNION BANK OF INDIA',
    'UTBI': 'UNITED BANK OF INDIA',
    'VIJB': 'VIJAYA BANK',
    'BMBL': 'BHARTIYA MAHILA BANK',
    
    # Private Sector Banks
    'UTIB': 'AXIS BANK',
    'HDFC': 'HDFC BANK',
    'ICIC': 'ICICI BANK',
    'INDB': 'INDUSIND BANK',
    'KKBK': 'KOTAK MAHINDRA BANK',
    'YESB': 'YES BANK',
    'DCBL': 'DCB BANK',
    'FDRL': 'FEDERAL BANK',
    'KARB': 'KARNATAKA BANK',
    'KVBL': 'KARUR VYSYA BANK',
    'RATN': 'RBL BANK',
    'SIBL': 'SOUTH INDIAN BANK',
    'TMBL': 'TAMILNAD MERCANTILE BANK',
    'VYSA': 'ING VYSYA BANK',
    
    # Co-operative Banks
    'ABHY': 'ABHYUDAYA CO-OP BANK',
    'ASBL': 'APNA SAHAKARI BANK',
    'GSCB': 'GUJARAT STATE CO-OP BANK',
    'HCBL': 'HASTI CO-OP BANK',
    'JSBP': 'JANATA SAHAKARI BANK',
    'MSNU': 'MEHSANA URBAN CO-OP BANK',
    'NTBL': 'NAINITAL BANK',
    'NKGS': 'NKGSB CO-OP BANK',
    'PMCB': 'PUNJAB & MAHARASHTRA CO-OP BANK',
    'SRCB': 'SARASWAT BANK',
    
    # Additional common banks
    'CIUB': 'CITY UNION BANK',
    'CSBK': 'CATHOLIC SYRIAN BANK',
    'DLXB': 'DHANLAXMI BANK',
    'ESFB': 'EQUITAS SMALL FINANCE BANK',
    'IDFB': 'IDFC FIRST BANK',
    'JAKA': 'JAMMU & KASHMIR BANK',
    'LAVB': 'LAKSHMI VILAS BANK',
    'NSPB': 'NSDL PAYMENTS BANK',
    'PAYU': 'PAYU PAYMENTS PRIVATE LIMITED',
    'PYTM': 'PAYTM PAYMENTS BANK',
}

def get_bank_name(ifsc_code):
    """
    Extract bank name from IFSC code
    IFSC Structure: XXXX0YYYYYY
    - First 4 characters (XXXX): Bank identifier
    - 5th character: Always '0'
    - Last 6 characters (YYYYYY): Branch code
    """
    if not ifsc_code or len(ifsc_code) < 4:
        return f"UNKNOWN BANK ({ifsc_code})"
    
    bank_code = ifsc_code[:4].upper()
    bank_name = IFSC_BANK_MAP.get(bank_code)
    
    if bank_name:
        return bank_name
    else:
        # Return a generic name with the bank code if not found in our map
        return f"{bank_code} BANK"

def validate_ifsc(ifsc_code):
    """
    Validate IFSC code format
    Should be 11 characters: 4 alphabetic + 0 + 6 alphanumeric
    """
    if len(ifsc_code) != 11:
        return False
    
    # First 4 should be alphabetic
    if not ifsc_code[:4].isalpha():
        return False
    
    # 5th should be 0
    if ifsc_code[4] != '0':
        return False
    
    # Last 6 can be alphanumeric
    if not ifsc_code[5:].isalnum():
        return False
    
    return True

def read_excel_data(excel_file):
    """Read and group data from Excel or CSV by IFSC code"""
    try:
        if excel_file.lower().endswith('.csv'):
            df = pd.read_csv(excel_file)
        else:
            df = pd.read_excel(excel_file)
        
        # Check required columns (case-insensitive)
        df.columns = df.columns.str.strip()
        
        # Try to find columns with flexible naming
        account_col = None
        name_col = None
        ifsc_col = None
        
        for col in df.columns:
            col_lower = col.lower()
            if (
                (
                    'account' in col_lower and (
                        'number' in col_lower or 'no' in col_lower or '#' in col_lower
                    )
                )
                or ('a/c' in col_lower or 'ac no' in col_lower or 'acc no' in col_lower or 'acno' in col_lower)
            ) and 'name' not in col_lower:
                account_col = account_col or col
            elif (
                ('account' in col_lower and 'name' in col_lower)
                or ('name' in col_lower and ('beneficiary' in col_lower or 'holder' in col_lower))
                or ('account' in col_lower and 'holder' in col_lower)
            ):
                name_col = name_col or col
            elif 'ifsc' in col_lower:
                ifsc_col = ifsc_col or col
        
        if not all([account_col, name_col, ifsc_col]):
            print("Error: Required columns not found in Excel file")
            print(f"Available columns: {', '.join(df.columns)}")
            print("\nRequired columns (flexible naming):")
            print("  - Account Number (or similar)")
            print("  - Account Name (or similar)")
            print("  - IFSC (or similar)")
            return None
        
        # Group by IFSC code
        grouped = defaultdict(list)
        invalid_ifsc = []
        
        for _, row in df.iterrows():
            ifsc = str(row[ifsc_col]).strip().upper()
            
            # Validate IFSC format
            if not validate_ifsc(ifsc):
                invalid_ifsc.append(ifsc)
                continue
            
            grouped[ifsc].append({
                'account_no': str(row[account_col]).strip(),
                'account_name': str(row[name_col]).strip(),
                'ifsc': ifsc
            })
        
        if invalid_ifsc:
            print(f"\n⚠ Warning: Found {len(invalid_ifsc)} invalid IFSC codes (skipped):")
            for ifsc in invalid_ifsc[:5]:  # Show first 5
                print(f"  - {ifsc}")
            if len(invalid_ifsc) > 5:
                print(f"  ... and {len(invalid_ifsc) - 5} more")
        
        return grouped
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return None

def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '8')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)

def _set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '8')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tcPr.append(borders)

def _detect_tone(doc):
    text = ' '.join(p.text.lower() for p in doc.paragraphs)
    if any(k in text for k in ['urgent', 'immediate', 'final notice', 'last reminder']):
        return 'urgent'
    if any(k in text for k in ['kindly', 'please', 'request', 'cooperate']):
        return 'friendly'
    return 'formal'

def _apply_tone(doc, tone, font_name_override=None, font_size_override=None):
    if tone == 'auto':
        tone = _detect_tone(doc)
    if tone not in ['formal', 'urgent', 'friendly']:
        tone = 'formal'
    font_name = font_name_override
    font_size = font_size_override
    for p in doc.paragraphs:
        if font_name and font_size:
            for r in p.runs:
                r.font.name = font_name
                r.font.size = Pt(font_size)
        if tone == 'urgent' and any(k in p.text.lower() for k in ['notice', 'urgent']):
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def _set_cell_margins(cell, top=36, left=36, bottom=36, right=36):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def _apply_paragraph_style(paragraphs, font_name, font_size, spacing):
    for p in paragraphs:
        for r in p.runs:
            r.font.name = font_name
            r.font.size = Pt(font_size)
        pf = p.paragraph_format
        if spacing is not None:
            sb, sa, ls = spacing
            if sb is not None:
                pf.space_before = sb
            if sa is not None:
                pf.space_after = sa
            if ls is not None:
                pf.line_spacing = ls

def _get_cell_width(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    w_el = tcPr.find(qn('w:tcW'))
    if w_el is not None and w_el.get(qn('w:w')):
        try:
            return int(w_el.get(qn('w:w')))
        except Exception:
            return None
    return None

def _set_cell_width(cell, width_twips):
    if width_twips is None:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    w_el = tcPr.find(qn('w:tcW'))
    if w_el is None:
        w_el = OxmlElement('w:tcW')
        tcPr.append(w_el)
    w_el.set(qn('w:type'), 'dxa')
    w_el.set(qn('w:w'), str(width_twips))

def _extract_template_baseline(doc, fallback_font_name='Bookman Old Style', fallback_font_size=8):
    font_name = None
    font_size = None
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font and r.font.name:
                font_name = r.font.name
            if r.font and r.font.size:
                try:
                    font_size = int(r.font.size.pt)
                except Exception:
                    pass
            if font_name and font_size:
                break
        if font_name and font_size:
            break
    if not font_name:
        font_name = fallback_font_name
    if not font_size:
        font_size = fallback_font_size
    header_widths = None
    header_row_height = None
    header_spacing = None
    for table in doc.tables:
        if len(table.columns) == 3:
            header_text = ''.join([cell.text for cell in table.rows[0].cells]).lower()
            if 'account' in header_text and 'ifsc' in header_text:
                header_widths = [_get_cell_width(c) for c in table.rows[0].cells]
                try:
                    header_row_height = table.rows[0].height
                except Exception:
                    header_row_height = None
                try:
                    hp = table.rows[0].cells[0].paragraphs[0]
                    pf = hp.paragraph_format
                    header_spacing = (pf.space_before, pf.space_after, pf.line_spacing)
                except Exception:
                    header_spacing = None
                break
    return font_name, font_size, header_widths, header_row_height, header_spacing

def _extract_nodal_baseline(doc, fallback_font, fallback_size):
    style_name = None
    alignment = None
    spacing = None
    indent = None
    font_name = None
    font_size = None
    bold = None
    for i, p in enumerate(doc.paragraphs):
        if 'nodal officer' in p.text.lower():
            if i + 1 < len(doc.paragraphs):
                np = doc.paragraphs[i + 1]
                try:
                    style_name = np.style.name if np.style else None
                except Exception:
                    style_name = None
                pf = np.paragraph_format
                alignment = pf.alignment
                spacing = (pf.space_before, pf.space_after, pf.line_spacing)
                indent = (pf.left_indent, pf.right_indent, pf.first_line_indent)
                for r in np.runs:
                    if r.font and r.font.name and not font_name:
                        font_name = r.font.name
                    if r.font and r.font.size and not font_size:
                        try:
                            font_size = int(r.font.size.pt)
                        except Exception:
                            pass
                    if bold is None and r.bold is not None:
                        bold = r.bold
                break
    if not font_name:
        font_name = fallback_font
    if not font_size:
        font_size = fallback_size
    return {
        'style_name': style_name,
        'alignment': alignment,
        'spacing': spacing,
        'indent': indent,
        'font_name': font_name,
        'font_size': font_size,
        'bold': bold,
    }

def _apply_baseline_to_paragraph(paragraph, text, baseline, doc):
    paragraph.text = text
    if baseline.get('style_name'):
        try:
            paragraph.style = doc.styles[baseline['style_name']]
        except Exception:
            pass
    pf = paragraph.paragraph_format
    if baseline.get('alignment') is not None:
        pf.alignment = baseline['alignment']
    spacing = baseline.get('spacing')
    if spacing:
        sb, sa, ls = spacing
        if sb is not None:
            pf.space_before = sb
        if sa is not None:
            pf.space_after = sa
        if ls is not None:
            pf.line_spacing = ls
    indent = baseline.get('indent')
    if indent:
        li, ri, fi = indent
        if li is not None:
            pf.left_indent = li
        if ri is not None:
            pf.right_indent = ri
        if fi is not None:
            pf.first_line_indent = fi
    for r in paragraph.runs:
        r.font.name = baseline.get('font_name')
        fs = baseline.get('font_size')
        if fs:
            r.font.size = Pt(fs)
        if baseline.get('bold') is not None:
            r.bold = baseline['bold']

def update_word_template(template_file, output_file, bank_name, accounts, placeholder='ICICI BANK', tone='formal', font_name='Bookman Old Style', font_size=8):
    """Update Word template with bank name and accounts table"""
    try:
        doc = Document(template_file)
        tmpl_font, tmpl_size, header_widths, header_row_height, header_spacing = _extract_template_baseline(doc, font_name, font_size)
        _apply_tone(doc, tone)  # do not override template fonts globally
        
        # Update bank name in the document
        # Search in paragraphs
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, bank_name)
        
        # Search in tables (headers, etc.)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, bank_name)
        
        # Find and update the accounts table (table with 3 columns)
        for table in doc.tables:
            if len(table.columns) == 3:
                header_text = ''.join([cell.text for cell in table.rows[0].cells]).lower()
                if 'account' in header_text and 'ifsc' in header_text:
                    header_row = table.rows[0]
                    for cell in header_row.cells:
                        _apply_paragraph_style(cell.paragraphs, tmpl_font, tmpl_size, header_spacing)
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.bold = True
                        _set_cell_margins(cell, 36, 36, 36, 36)
                    header_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                    header_row.height = header_row_height or Pt(tmpl_size + 6)
                    rows_to_delete = len(table.rows) - 1
                    for _ in range(rows_to_delete):
                        table._element.remove(table.rows[-1]._element)
                    for account in accounts:
                        row = table.add_row()
                        row.cells[0].text = account['account_no']
                        row.cells[1].text = account['account_name']
                        row.cells[2].text = account['ifsc']
                        for cell in row.cells:
                            _apply_paragraph_style(cell.paragraphs, tmpl_font, tmpl_size, header_spacing)
                            _set_cell_borders(cell)
                            _set_cell_margins(cell, 36, 36, 36, 36)
                        if header_widths:
                            for idx, cell in enumerate(row.cells):
                                _set_cell_width(cell, header_widths[idx] if idx < len(header_widths) else None)
                        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                        row.height = header_row_height or Pt(tmpl_size + 6)
                    try:
                        table.style = 'Table Grid'
                    except Exception:
                        pass
                    _set_table_borders(table)
                    break

        nodal_baseline = _extract_nodal_baseline(doc, tmpl_font, tmpl_size)
        for i, paragraph in enumerate(doc.paragraphs):
            if 'nodal officer' in paragraph.text.lower():
                if i + 1 < len(doc.paragraphs):
                    next_p = doc.paragraphs[i + 1]
                    _apply_baseline_to_paragraph(next_p, bank_name, nodal_baseline, doc)

        for table in doc.tables:
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if 'nodal officer' in cell.text.lower():
                        if r_idx + 1 < len(table.rows):
                            target_cell = table.cell(r_idx + 1, c_idx)
                            for p in target_cell.paragraphs:
                                _apply_baseline_to_paragraph(p, bank_name, nodal_baseline, doc)
        
        # Save the document
        doc.save(output_file)
        return True
    except Exception as e:
        print(f"Error updating Word template: {e}")
        return False

def main():
    parser = argparse.ArgumentParser(
        description='Generate bank notices from Excel file and Word template',
        epilog='Example: python notice_generator.py accounts.xlsx template.docx -o output_notices'
    )
    parser.add_argument('excel_file', help='Path to Excel file with account data')
    parser.add_argument('template_file', help='Path to Word template file')
    parser.add_argument('-o', '--output-dir', default='notices_output',
                       help='Output directory for generated notices (default: notices_output)')
    parser.add_argument('--placeholder', default='ICICI BANK',
                       help='Placeholder text in template to replace with bank name (default: ICICI BANK)')
    parser.add_argument('--tone', default='formal',
                       help='Document tone: formal, urgent, friendly, or auto (default: formal)')
    parser.add_argument('--font-name', default='Bookman Old Style',
                       help='Default font name for document and tables (default: Bookman Old Style)')
    parser.add_argument('--font-size', type=int, default=8,
                       help='Default font size in points (default: 8)')
    
    args = parser.parse_args()
    
    # Check if files exist
    if not os.path.exists(args.excel_file):
        print(f"❌ Error: Excel file '{args.excel_file}' not found")
        return
    
    if not os.path.exists(args.template_file):
        print(f"❌ Error: Template file '{args.template_file}' not found")
        return
    
    # Create output directory
    os.makedirs(args.output_dir, exist_ok=True)
    
    print(f"📖 Reading data from {args.excel_file}...")
    grouped_data = read_excel_data(args.excel_file)
    
    if not grouped_data:
        print("❌ Failed to read Excel data")
        return
    
    print(f"\n✅ Found {len(grouped_data)} unique IFSC codes")
    total_accounts = sum(len(accounts) for accounts in grouped_data.values())
    print(f"📊 Total accounts: {total_accounts}")
    print(f"\n🔨 Generating notices...\n")
    
    # Generate notice for each IFSC group
    success_count = 0
    for ifsc, accounts in grouped_data.items():
        bank_name = get_bank_name(ifsc)
        output_file = os.path.join(args.output_dir, f"Notice_{bank_name.replace(' ', '_')}_{ifsc}.docx")
        
        print(f"  📄 {bank_name} ({ifsc}) - {len(accounts)} account(s)")
        
        if update_word_template(args.template_file, output_file, bank_name, accounts, args.placeholder, args.tone, args.font_name, args.font_size):
            print(f"     ✅ Saved: {output_file}")
            success_count += 1
        else:
            print(f"     ❌ Failed")
    
    print(f"\n{'='*60}")
    print(f"✅ Done! Generated {success_count}/{len(grouped_data)} notices")
    print(f"📁 All notices saved in '{args.output_dir}' directory")
    print(f"{'='*60}")

if __name__ == '__main__':
    main()